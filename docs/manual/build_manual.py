"""
Build the Pharma Glimmora — AI Backend Integration user manual (.docx).

Run from the project root:
    python docs/manual/build_manual.py

Outputs docs/manual/Pharma-Glimmora-AI-Manual.docx with embedded screenshots.

The screenshots in docs/manual/screenshots/ are captured separately via
Playwright; this script just assembles them with structured prose,
endpoint reference tables, and JSON sample blocks.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx is required. Install with: pip install python-docx\n")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUT = ROOT / "Pharma-Glimmora-AI-Manual.docx"

BRAND_AMBER = RGBColor(0x8B, 0x69, 0x14)
BRAND_DARK = RGBColor(0x30, 0x2D, 0x29)
GRAY = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xDC, 0x26, 0x26)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_DARK


def add_para(doc: Document, text: str, *, bold: bool = False, color: RGBColor | None = None, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_image(doc: Document, filename: str, caption: str | None = None, width: float = 6.5) -> None:
    img_path = SHOTS / filename
    if not img_path.exists():
        add_para(doc, f"[missing screenshot: {filename}]", italic=True, color=GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"Figure — {caption}")
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = GRAY


def add_code_block(doc: Document, text: str) -> None:
    """Mono-style multi-line block."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Force the code font on East-Asian fallback too.
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)


def add_endpoint_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    """Three-column table: Method+Path, UI surface, Sample call."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Endpoint", "Where in app", "Notes")):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = BRAND_AMBER
        set_cell_shading(hdr[i], "FFF7E6")
    for i, (a, b, c) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pharma Glimmora")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_AMBER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI Backend Integration — User Manual")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Backend: https://pharma-glimmora-ai-backend.onrender.com\n"
        "Frontend: Pharma Glimmora — GxP Compliance Platform"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_paragraph()
    add_para(
        doc,
        "This manual covers every endpoint exposed by the AI backend, "
        "where each one surfaces inside the Pharma Glimmora frontend, and "
        "sample inputs/outputs. Screenshots show the UI surface for each "
        "feature.",
        size=11,
    )

    doc.add_page_break()


def section_overview(doc: Document) -> None:
    add_heading(doc, "1. Overview", level=1)
    add_para(
        doc,
        "The AI backend is a FastAPI service that exposes endpoints for "
        "authentication, AI-assisted CAPA generation, the full CAPA "
        "lifecycle (RCA, Action Plan, Implementation Monitoring, "
        "Effectiveness Check, Closure), an audit trail, and a "
        "conversational AI assistant with optional voice round-trip.",
    )
    add_para(
        doc,
        "All protected endpoints require an access token in an `auth` "
        "request header. The token is obtained from /api/v1/auth/login "
        "and is automatically refreshed by the frontend on every app "
        "sign-in — no separate sign-in for the AI features.",
    )

    add_heading(doc, "Endpoint inventory", level=2)
    rows = [
        ("POST /api/v1/auth/login", "Every app login (LoginPage refreshAiToken)", "JSON {username, password} → {access_token, …}"),
        ("POST /api/v1/auth/signup", "Customer admin / tenant user create + login self-heal retry", "Auto-runs on first login if user not registered yet"),
        ("POST /api/ai/chat", "Floating chatbot — text mode", "Threads chat_history; supports follow-up context"),
        ("POST /api/ai/voice/chat", "Floating chatbot — round-trip mic", "audio in → audio reply (Whisper + chat + TTS)"),
        ("POST /api/ai/voice/transcribe", "Chatbot — Dictate button", "STT only; transcript drops into the input"),
        ("POST /api/ai/voice/speak", "Speaker icon on every assistant text bubble", "TTS only; plays one stored reply on demand"),
        ("POST /api/v1/capa/create", "AI CAPA modal (CAPA Tracker)", "Multipart form with optional document; returns risk score + similar CAPAs + recommendation"),
        ("GET /api/v1/capa/all", "/ai-capa index — super_admin only", "Returns every CAPA across customers"),
        ("GET /api/v1/capa/customer/{id}", "/ai-capa index — default scope", "Customer-scoped CAPA list"),
        ("GET /api/v1/capa/status/{id}", "/ai-capa/[id] — Summary card", "Full record by capa_id"),
        ("POST /api/v1/capa/dismiss-alert", "Recurrence panel inside the AI CAPA result modal", "Reason ≥ 5 chars + e-signature, immutable in audit trail"),
        ("POST /api/v1/rca/submit", "RCA modal in lifecycle page", "Method (5-Why / Fishbone / Fault Tree / Other) + evidence"),
        ("GET /api/v1/rca/capa/{id}", "RCA section in lifecycle page", "Auto-loaded; shows RCA records for the CAPA"),
        ("GET /api/v1/rca/status/{id}", "/ai-tools — RCA status lookup", "Fetch a single RCA by rca_id"),
        ("POST /api/v1/action-plan/submit", "Action Plan modal", "Multi-row action editor (description / responsible / due date)"),
        ("GET /api/v1/action-plan/capa/{id}", "Action Plan section in lifecycle page", "Auto-loaded"),
        ("GET /api/v1/action-plan/status/{id}", "/ai-tools — Action plan status lookup", "Fetch a single plan by action_plan_id"),
        ("POST /api/v1/monitoring/check", "Monitoring modal", "Per-action status (On Track / Delayed / …) + progress notes"),
        ("GET /api/v1/monitoring/capa/{id}", "Monitoring section in lifecycle page", "Auto-loaded"),
        ("GET /api/v1/monitoring/status/{id}", "/ai-tools — Monitoring status lookup", ""),
        ("POST /api/v1/effectiveness/check", "Effectiveness modal", "Days since CAPA, evidence items, trend metrics, new-issues flag"),
        ("GET /api/v1/effectiveness/capa/{id}", "Effectiveness section in lifecycle page", "Auto-loaded"),
        ("GET /api/v1/effectiveness/status/{id}", "/ai-tools — Effectiveness status lookup", ""),
        ("POST /api/v1/closure/initiate", "Closure modal", "21 CFR Part 11: approver, designation, e-signature, rationale"),
        ("GET /api/v1/closure/capa/{id}", "Closure section in lifecycle page", "Auto-loaded"),
        ("GET /api/v1/closure/status/{id}", "/ai-tools — Closure status lookup", ""),
        ("GET /api/v1/audit/all", "Audit Trail page → AI Backend tab", "Returns the full backend audit log"),
        ("GET /api/v1/audit/record/{id}", "/ai-tools — Audit record lookup", "Fetch one audit entry by audit_id"),
        ("GET /api/v1/users/", "/ai-tools — Users endpoint diagnostic", "Returns 'Users endpoint working ✅'"),
        ("GET /api/ai/health", "/ai-tools — AI Assistant health", "Returns service status"),
        ("GET /api/ai/voice/health", "/ai-tools — AI Voice health", "Returns voice status + supported voices"),
    ]
    add_endpoint_table(doc, rows)
    doc.add_page_break()


def section_login(doc: Document) -> None:
    add_heading(doc, "2. Sign in", level=1)
    add_para(
        doc,
        "The login page accepts both usernames and email addresses. On "
        "success, the frontend silently calls /api/v1/auth/login on the "
        "AI backend and stores the returned access_token on the user "
        "record. If the user isn't registered on the AI backend yet, the "
        "frontend auto-runs /api/v1/auth/signup with the credentials "
        "you typed — no manual provisioning step is required.",
    )
    add_image(doc, "01-login.png", caption="Sign-in page")
    add_heading(doc, "Sample login response", level=3)
    add_code_block(
        doc,
        "POST /api/v1/auth/login\n"
        '{"username":"qa@pharmaglimmora.com","password":"QaHead@123"}\n\n'
        "200 OK\n"
        "{\n"
        '  "access_token": "eyJhbGciOiJIUzI1NiIs…",\n'
        '  "token_type": "Bearer",\n'
        '  "username": "qa@pharmaglimmora.com",\n'
        '  "customer_id": "CUST_TEST01",\n'
        '  "role": "qa_head",\n'
        '  "message": "✅ Login successful. Welcome qa@pharmaglimmora.com!"\n'
        "}",
    )
    doc.add_page_break()


def section_capa(doc: Document) -> None:
    add_heading(doc, "3. AI-Generated CAPA", level=1)
    add_para(
        doc,
        "On the CAPA Tracker page, click the AI CAPA button (next to "
        "New CAPA) to open the AI-Generated CAPA modal. Customer ID is "
        "auto-populated from your tenant's customer admin and is no "
        "longer shown in the form.",
    )
    add_image(doc, "03-capa-tracker.png", caption="CAPA Tracker — AI CAPA + New CAPA buttons")
    add_image(doc, "04-ai-capa-modal.png", caption="AI-Generated CAPA modal")
    add_heading(doc, "Sample input (multipart/form-data)", level=3)
    add_code_block(
        doc,
        "POST /api/v1/capa/create\n"
        "Headers: auth: <access_token>\n\n"
        "  customer_id      = CUST_001\n"
        "  problem_statement = Coating defects observed on tablet batch\n"
        "                      BX-2026-088 — inconsistent film thickness\n"
        "                      on ~12% of units.\n"
        "  source           = Deviation\n"
        "  area_affected    = Manufacturing - Coating Suite 2\n"
        "  equipment_product = Coater (GLATT GC-1000)\n"
        "  initial_severity = High\n"
        "  document         = (optional file upload)\n",
    )
    add_heading(doc, "Sample response", level=3)
    add_code_block(
        doc,
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "status": "Open",\n'
        '  "created_at": "2026-04-30T12:30:26.922497",\n'
        '  "is_recurring": true,\n'
        '  "similar_capas": [\n'
        "    {\n"
        '      "capa_id": "CAPA-2023-089",\n'
        '      "similarity_score": 0.95,\n'
        '      "description": "Both involve coating defects on similar equipment.",\n'
        '      "was_effective": false\n'
        "    }\n"
        "  ],\n"
        '  "recurrence_alert": "Recurring coating defects detected, with prior ineffective CAPA.",\n'
        '  "pattern_detected": "Coating defects on coater equipment are recurring.",\n'
        '  "ai_recommendation": "Investigate root cause of coating defects on coater equipment.",\n'
        '  "risk_score": 0.85,\n'
        '  "message": "⚠️ RECURRING ISSUE — CAPA-2026-304 created. Escalation recommended."\n'
        "}",
    )
    add_heading(doc, "Sample inputs that produce rich responses", level=3)
    add_para(
        doc,
        "These are validated examples; the first reliably triggers the "
        "recurring-issue path:",
        italic=True,
    )
    samples = [
        ("Coating defects (recurring)",
         "Source: Deviation · Area: Manufacturing - Coating Suite 2 · Equipment: Coater (GLATT GC-1000) · Severity: High\n"
         "Problem: Film coating thickness variation observed on tablet batch BX-2026-088. ~12% of units below spec at QC. "
         "Operator reports erratic spray pattern from gun #3."),
        ("Tablet hardness OOS",
         "Source: OOS · Area: Manufacturing - Compression Suite 1 · Equipment: Tablet Press (Korsch XL-100) · Severity: Critical\n"
         "Problem: Tablet hardness OOS — average 4.2 kP vs spec 6–10 kP on stratified sampling of batch TX-2026-044. "
         "Compression force log shows downward drift over 2 hrs."),
        ("Cleaning validation failure",
         "Source: Deviation · Area: API Manufacturing - Reactor Bay · Equipment: Reactor R-301 (Pfaudler 500L) · Severity: Critical\n"
         "Problem: Swab samples from CIP cycle on Reactor R-301 returned TOC 18 ppb against 10 ppb residue limit. "
         "Adjacent product changeover scheduled for tomorrow."),
        ("HVAC excursion",
         "Source: Deviation · Area: Sterile Filling - ISO 8 Cleanroom · Equipment: HVAC AHU-04 · Severity: Critical\n"
         "Problem: Class 100,000 cleanroom particle count exceeded action limit (0.5µm > 3,520,000/m³) for 14 minutes "
         "during aseptic fill of batch SF-2026-012."),
    ]
    for title, body in samples:
        add_para(doc, "• " + title, bold=True, size=10, color=BRAND_AMBER)
        add_para(doc, body, size=10)
    add_para(
        doc,
        "On Accept & close, the user is navigated to the AI CAPA "
        "Lifecycle page (/ai-capa/[capaId]) so they can run the rest "
        "of the lifecycle.",
        size=10,
    )
    doc.add_page_break()


def section_lifecycle(doc: Document) -> None:
    add_heading(doc, "4. AI CAPA Lifecycle", level=1)
    add_para(
        doc,
        "The lifecycle page at /ai-capa/[capaId] is the operational "
        "centre for one CAPA. It loads every stage in parallel and "
        "presents each one as a card with either the data or a Submit "
        "button. Stages are gated — Action Plan unlocks once an RCA "
        "exists, Monitoring and Effectiveness unlock once the Action "
        "Plan exists, and Closure requires Effectiveness.",
    )
    add_image(doc, "06-ai-capa-lifecycle.png", caption="AI CAPA Lifecycle dashboard")

    add_heading(doc, "RCA submission", level=2)
    add_code_block(
        doc,
        "POST /api/v1/rca/submit\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "rca_method": "5-Why",\n'
        '  "evidence": "Operator log shows spray nozzle pressure drift over 2 hours."\n'
        "}",
    )

    add_heading(doc, "Action Plan submission", level=2)
    add_code_block(
        doc,
        "POST /api/v1/action-plan/submit\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "rca_id": "RCA-2026-101",\n'
        '  "actions": [\n'
        "    {\n"
        '      "action_description": "Recalibrate spray nozzle gun #3 and update SOP",\n'
        '      "responsible_person": "QA Manager",\n'
        '      "due_date": "2026-05-15"\n'
        "    },\n"
        "    {\n"
        '      "action_description": "Retrain all operators on coater calibration",\n'
        '      "responsible_person": "Training Lead",\n'
        '      "due_date": "2026-05-30"\n'
        "    }\n"
        "  ]\n"
        "}",
    )

    add_heading(doc, "Monitoring check", level=2)
    add_code_block(
        doc,
        "POST /api/v1/monitoring/check\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "action_plan_id": "AP-2026-201",\n'
        '  "action_updates": [\n'
        "    {\n"
        '      "action_description": "Recalibrate spray nozzle gun #3",\n'
        '      "responsible_person": "QA Manager",\n'
        '      "due_date": "2026-05-15",\n'
        '      "status": "Completed",\n'
        '      "progress_note": "Calibrated 2026-05-12; new SOP issued"\n'
        "    }\n"
        "  ]\n"
        "}",
    )

    add_heading(doc, "Effectiveness check", level=2)
    add_code_block(
        doc,
        "POST /api/v1/effectiveness/check\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "action_plan_id": "AP-2026-201",\n'
        '  "days_since_capa": 90,\n'
        '  "evidence_items": [\n'
        "    {\n"
        '      "action_description": "Recalibrate spray nozzle gun #3",\n'
        '      "completed": true,\n'
        '      "evidence_attached": true,\n'
        '      "evidence_note": "Calibration record + updated SOP attached"\n'
        "    }\n"
        "  ],\n"
        '  "trend_data": [\n'
        "    {\n"
        '      "metric_name": "Coating defects",\n'
        '      "before_capa": 4.2,\n'
        '      "after_capa": 0.5,\n'
        '      "unit": "per month"\n'
        "    }\n"
        "  ],\n"
        '  "new_issues_reported": false\n'
        "}",
    )

    add_heading(doc, "Closure", level=2)
    add_para(
        doc,
        "Closure is gated by Effectiveness and requires a 21 CFR Part 11 "
        "compliant electronic signature. The dialog warns that the "
        "signature is logged immutably.",
    )
    add_code_block(
        doc,
        "POST /api/v1/closure/initiate\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "customer_id": "CUST_001",\n'
        '  "effectiveness_id": "EFF-2026-501",\n'
        '  "approved_by": "Dr. Priya Sharma",\n'
        '  "designation": "QA Head",\n'
        '  "electronic_signature": "PS-2026-QA",\n'
        '  "closure_rationale": "All actions completed with evidence. No recurrence in 90 days.",\n'
        '  "related_capas_reviewed": true,\n'
        '  "document_changes_approved": true\n'
        "}",
    )
    doc.add_page_break()


def section_dismiss_alert(doc: Document) -> None:
    add_heading(doc, "5. Dismissing recurrence alerts", level=1)
    add_para(
        doc,
        "When the AI CAPA result panel surfaces a recurrence alert, a "
        "Dismiss alert button sits directly under it. Clicking expands "
        "an inline form with a reason field (≥ 5 characters) and an "
        "electronic signature field. The dismissal is logged in the "
        "audit trail with dismissed_by set to the logged-in user's name.",
    )
    add_code_block(
        doc,
        "POST /api/v1/capa/dismiss-alert\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "capa_id": "CAPA-2026-304",\n'
        '  "alert_type": "recurrence_alert",\n'
        '  "dismissal_reason": "Investigated and confirmed different root cause; ' "see RCA-2026-101.\",\n"
        '  "electronic_signature": "PS-2026-001",\n'
        '  "dismissed_by": "Dr. Priya Sharma"\n'
        "}",
    )
    doc.add_page_break()


def section_index(doc: Document) -> None:
    add_heading(doc, "6. AI CAPA index", level=1)
    add_para(
        doc,
        "The /ai-capa page lists every backend CAPA for the current "
        "customer, with quick stats (total / open / recurring / high "
        "risk), per-row severity / status / risk-score badges, and "
        "click-through to the lifecycle page. Super admins also get a "
        "My customer / All customers toggle.",
    )
    add_image(doc, "05-ai-capa-index.png", caption="AI CAPAs index")
    add_heading(doc, "Sample response (capa/customer)", level=3)
    add_code_block(
        doc,
        "GET /api/v1/capa/customer/CUST_001\n"
        "Headers: auth: <access_token>\n\n"
        "200 OK\n"
        "{\n"
        '  "customer_id": "CUST_001",\n'
        '  "total": 4,\n'
        '  "capas": [\n'
        "    {\n"
        '      "capa_id": "CAPA-2026-304",\n'
        '      "problem_statement": "tablet coating",\n'
        '      "source": "deviation",\n'
        '      "severity": "major",\n'
        '      "status": "Open",\n'
        '      "is_recurring": true,\n'
        '      "risk_score": 0.85,\n'
        '      "created_at": "2026-04-30T12:30:26.922497"\n'
        "    }\n"
        "  ]\n"
        "}",
    )
    doc.add_page_break()


def section_audit(doc: Document) -> None:
    add_heading(doc, "7. Audit Trail", level=1)
    add_para(
        doc,
        "The Audit Trail page sits in the Readiness & Governance "
        "section. The Local tab shows the in-app Redux audit log; the "
        "AI Backend tab fetches /api/v1/audit/all and renders the "
        "backend's audit_logs collection (action_type, feature_id, "
        "record_id, username, status, audit_id, timestamp).",
    )
    add_image(doc, "08-audit-local.png", caption="Audit Trail — Local tab")
    add_image(doc, "09-audit-ai.png", caption="Audit Trail — AI Backend tab")
    add_heading(doc, "Sample response", level=3)
    add_code_block(
        doc,
        "GET /api/v1/audit/all\n"
        "Headers: auth: <access_token>\n\n"
        "200 OK\n"
        "{\n"
        '  "total": 6,\n'
        '  "audit_logs": [\n'
        "    {\n"
        '      "audit_id": "AUDIT-20260430123026-ac1bd3b9",\n'
        '      "action_type": "create_capa",\n'
        '      "feature_id": "AI-100",\n'
        '      "record_id": "CAPA-2026-304",\n'
        '      "username": "qa_manager",\n'
        '      "status": "success",\n'
        '      "timestamp": "2026-04-30T12:30:26.951557"\n'
        "    }\n"
        "  ]\n"
        "}",
    )
    doc.add_page_break()


def section_tools(doc: Document) -> None:
    add_heading(doc, "8. AI Backend Tools", level=1)
    add_para(
        doc,
        "The /ai-tools page exposes every endpoint that doesn't have a "
        "dedicated UI surface — mostly the *Status/{id} GETs (which are "
        "redundant with the *ByCapa GETs already on the lifecycle "
        "page), plus audit/record, users/, and the two health checks. "
        "Each card mirrors the AI-Generated CAPA modal pattern: a "
        "header, a single input, a Submit button, and a JSON result "
        "panel.",
    )
    add_image(doc, "07-ai-tools.png", caption="AI Backend Tools page")
    add_heading(doc, "Sample diagnostic responses", level=3)
    add_code_block(
        doc,
        'GET /api/ai/health        → {"status": "AI assistant is running ✅"}\n'
        "GET /api/ai/voice/health  → {\n"
        '  "status": "Voice endpoints running ✅",\n'
        '  "endpoints": {"transcribe": "POST /api/ai/voice/transcribe",\n'
        '                "speak":      "POST /api/ai/voice/speak",\n'
        '                "voice_chat": "POST /api/ai/voice/chat"},\n'
        '  "voices": ["alloy","echo","fable","onyx","nova","shimmer"]\n'
        "}\n"
        'GET /api/v1/users/         → {"message": "Users endpoint working ✅"}\n',
    )
    doc.add_page_break()


def section_chatbot(doc: Document) -> None:
    add_heading(doc, "9. Floating AI Chatbot", level=1)
    add_para(
        doc,
        "Every authenticated page (except /login and /site-picker) "
        "carries a floating chatbot bubble bottom-right. Read-only "
        "viewers don't see it. Right-click and drag the bubble to move "
        "it; position is cached per browser. Click to open the panel.",
    )
    add_image(doc, "10-chatbot-open.png", caption="Floating chatbot — empty panel")

    add_heading(doc, "Text mode", level=2)
    add_para(
        doc,
        "Type into the input box and press Enter to send. The full "
        "running history is sent with each message so the assistant "
        "answers in context.",
    )
    add_code_block(
        doc,
        "POST /api/ai/chat\n"
        "Headers: auth: <access_token>\n\n"
        "{\n"
        '  "message": "How many CAPAs are open right now?",\n'
        '  "chat_history": [\n'
        '    {"role": "user", "content": "Tell me about CAPAs"},\n'
        '    {"role": "assistant", "content": "A CAPA is …"}\n'
        "  ]\n"
        "}\n\n"
        "200 OK\n"
        "{\n"
        '  "reply": "There are currently 3 open CAPAs for CUST_001 …",\n'
        '  "intent": "DB_QUERY",\n'
        '  "customer_id": "CUST_001"\n'
        "}",
    )

    add_heading(doc, "Voice round-trip (mic button)", level=2)
    add_para(
        doc,
        "Click the mic button to record. The signal flows through a "
        "real-time DSP cleanup chain (high-pass → RNNoise → compressor → "
        "make-up gain) before it reaches Whisper, so background noise is "
        "suppressed at speech-recognition quality. A live VU-meter shows "
        "the audio level; the recording ends when you click Stop. After "
        "the preview, click Send to upload to /api/ai/voice/chat — the "
        "backend transcribes, generates a reply, and returns audio that "
        "auto-plays. The transcript and assistant text are also rendered "
        "in the chat for review.",
    )
    add_para(
        doc,
        "The gear icon in the chatbot header opens a settings drawer "
        "with a Noise suppression slider (0% = raw mic, 100% = full "
        "RNNoise wet signal). Setting persists across reloads and "
        "live-updates while you record.",
    )

    add_heading(doc, "Dictate (Edit pencil button)", level=2)
    add_para(
        doc,
        "The pencil icon next to the mic uses /api/ai/voice/transcribe "
        "to drop your speech into the input box as text. Useful when "
        "you want to review and edit before sending.",
    )

    add_heading(doc, "Speak this reply (speaker icon on assistant bubble)", level=2)
    add_para(
        doc,
        "Every assistant text message has a small speaker icon. Click "
        "to call /api/ai/voice/speak with voice=\"nova\" and play the "
        "audio.",
    )
    doc.add_page_break()


def section_signup(doc: Document) -> None:
    add_heading(doc, "10. User signup flows", level=1)
    add_para(
        doc,
        "Three places in the app create users on the AI backend, all "
        "via /api/v1/auth/signup:",
    )
    add_para(doc, "1. Customer Accounts (super admin) — when a new customer is created, the customer admin is registered with a generated CUST_xxx that becomes both customer_id and user_id.", size=10)
    add_para(doc, "2. Settings → Users (customer admin) — when a tenant user is added, signup runs with the user's id as user_id and the customer admin's CUST_xxx as customer_id.", size=10)
    add_para(doc, "3. Login self-heal — if the user logs in but isn't registered yet (e.g. seeded mock accounts), refreshAiToken auto-runs signup with the credentials they typed.", size=10)
    add_heading(doc, "Sample signup payload", level=3)
    add_code_block(
        doc,
        "POST /api/v1/auth/signup\n"
        "Content-Type: application/json\n\n"
        "{\n"
        '  "user_id": "USER-A1B2C3D4",\n'
        '  "username": "ramu",\n'
        '  "email": "ramu@example.com",\n'
        '  "password": "secret",\n'
        '  "customer_id": "CUST_001",\n'
        '  "role": "qa_manager"\n'
        "}\n\n"
        "201 Created\n"
        "{\n"
        '  "access_token": "eyJ…",\n'
        '  "token_type": "Bearer",\n'
        '  "username": "ramu",\n'
        '  "customer_id": "CUST_001",\n'
        '  "role": "qa_manager",\n'
        '  "message": "✅ Account created for ramu."\n'
        "}",
    )
    add_para(
        doc,
        "Validation rules (backend): username ≥ 1 char, password ≥ 1 "
        "char, email ≥ 5 chars. user_id and customer_id are opaque "
        "strings — the frontend generates user_id values on demand.",
        size=10,
        italic=True,
    )
    doc.add_page_break()


def section_quickref(doc: Document) -> None:
    add_heading(doc, "Appendix A — Quick reference", level=1)
    add_heading(doc, "Roles and visibility", level=2)
    add_para(doc, "• super_admin / customer_admin / qa_head / qc_lab_director / regulatory_affairs / csv_val_lead / it_cdo / operations_head / viewer", size=10)
    add_para(doc, "• Floating chatbot: visible to all roles except viewer.", size=10)
    add_para(doc, "• AI CAPA button on Tracker: customer_admin and any role with create permission.", size=10)
    add_para(doc, "• /ai-capa, /ai-capa/[id], /ai-tools, /audit-trail: every authenticated user.", size=10)

    add_heading(doc, "Common error responses", level=2)
    add_code_block(
        doc,
        "401 Invalid username or password.        wrong creds on /auth/login\n"
        "401 Invalid token. Please login.         missing or expired auth header\n"
        "422 Field required …                     missing form/body field\n"
        "404 (no body)                            not-yet-submitted stage on a *ByCapa GET\n"
        "500 AI chat failed: <ExceptionType>      uncaught error inside chat() / RAG / DB\n",
    )

    add_heading(doc, "Console logging", level=2)
    add_para(
        doc,
        "Every API call from the frontend logs to the browser console "
        "in a consistent format:",
        size=10,
    )
    add_code_block(
        doc,
        "[aiAuth]    POST /api/v1/auth/login → sending {username: …, password: ***}\n"
        "[aiAuth]    POST /api/v1/auth/login ✓ 200 (1841ms) {access_token: …}\n"
        "[aiBackend] POST /api/v1/capa/create → sending\n"
        "[aiBackend] POST /api/v1/capa/create ✓ 201 (1142ms)\n"
        "[aiChat]    POST /api/ai/chat → sending\n"
        "[aiChat]    POST /api/ai/chat ✓ 200 (4126ms)",
    )


def main() -> None:
    if not SHOTS.exists():
        sys.stderr.write(f"Screenshot folder not found: {SHOTS}\n")
        sys.exit(1)

    doc = Document()
    # Default font tweaks.
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    title_page(doc)
    section_overview(doc)
    section_login(doc)
    section_capa(doc)
    section_lifecycle(doc)
    section_dismiss_alert(doc)
    section_index(doc)
    section_audit(doc)
    section_tools(doc)
    section_chatbot(doc)
    section_signup(doc)
    section_quickref(doc)

    doc.save(OUT)
    print(f"Wrote {OUT} ({os.path.getsize(OUT) // 1024} KB)")


if __name__ == "__main__":
    main()
